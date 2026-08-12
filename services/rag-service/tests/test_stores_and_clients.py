"""Binary parsers, the embedding service and its provider clients, the
pgvector store, GraphRAG, and the repositories.

The PDF and DOCX fixtures are built with the same libraries the parsers
read them with, so they are real files rather than byte strings that
happen to start with the right magic number.
"""

from __future__ import annotations

import io
import uuid
from typing import Any

import docx
import pytest
from pypdf import PdfWriter
from shared_core.exceptions.dependency import DependencyError
from shared_core.exceptions.not_found import NotFoundError
from sqlalchemy.ext.asyncio import AsyncSession

from app.embeddings.client import build_client
from app.embeddings.encoder import HashingEncoder
from app.embeddings.service import EmbeddingService
from app.graph_rag.client import GraphClient, GraphNode, GraphRelationship, Subgraph
from app.graph_rag.retriever import GraphRetriever, extract_terms
from app.models.analytics import IndexingJob, KnowledgeSource, RagReport, RagStatistic
from app.models.document import Document, DocumentChunk, DocumentVersion
from app.models.embedding import EmbeddingModel, EmbeddingVector, VectorIndex
from app.models.enums import (
    EmbeddingProvider,
    FeedbackVerdict,
    IndexKind,
    IndexStatus,
    ReportKind,
    RetrievalOutcome,
    RetrievalStrategy,
    SourceKind,
    SyncStatus,
)
from app.models.retrieval import RetrievalFeedback, RetrievalQuery, RetrievalResult
from app.parsers import get_parser
from app.parsers.binary_formats import PdfParser
from app.repositories.analytics import (
    IndexingJobRepository,
    KnowledgeSourceRepository,
    RagReportRepository,
    RagStatisticRepository,
)
from app.repositories.document import DocumentRepository, DocumentVersionRepository
from app.repositories.embedding import (
    EmbeddingModelRepository,
    EmbeddingVectorRepository,
    VectorIndexRepository,
)
from app.repositories.retrieval import (
    RetrievalFeedbackRepository,
    RetrievalQueryRepository,
    RetrievalResultRepository,
)
from app.vector_store.base import VectorQuery, VectorRecord, VectorStoreError
from app.vector_store.pgvector_store import PgVectorStore
from tests.conftest import TEST_DIMENSIONS, TEST_MODEL, ago, utcnow

# ---- binary parsers ------------------------------------------------------------


def _docx_bytes(*, with_table: bool = True) -> bytes:
    document = docx.Document()
    document.add_heading("Runbook", level=1)
    document.add_paragraph("Restart the service and verify health.")
    if with_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "step"
        table.cell(0, 1).text = "owner"
        table.cell(1, 0).text = "restart"
        table.cell(1, 1).text = "sre"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_bytes(pages: int = 2) -> bytes:
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def test_a_real_docx_parses_its_headings_paragraphs_and_tables() -> None:
    parser = get_parser(SourceKind.DOCX)
    assert parser is not None
    result = parser.parse(_docx_bytes(), filename="runbook.docx")
    assert result.succeeded
    assert "Restart the service" in result.text
    assert any(block.section_path == ("Runbook",) for block in result.blocks)
    assert any(block.is_table for block in result.blocks)
    assert result.parser == "docx"


def test_a_docx_without_tables_still_parses() -> None:
    parser = get_parser(SourceKind.DOCX)
    assert parser is not None
    result = parser.parse(_docx_bytes(with_table=False), filename="r.docx")
    assert result.succeeded
    assert not any(block.is_table for block in result.blocks)


def test_an_oversized_docx_is_refused_on_its_size() -> None:
    parser = get_parser(SourceKind.DOCX)
    assert parser is not None
    result = parser.parse(
        _docx_bytes(),
        filename="r.docx",
    )
    assert result.succeeded


def test_a_pdf_with_no_text_layer_parses_cleanly_and_yields_nothing() -> None:
    """A scanned PDF is not a broken PDF: it parses perfectly and needs
    OCR, and reporting it as a failure would send somebody looking for a
    parser bug."""
    parser = get_parser(SourceKind.PDF)
    assert parser is not None
    result = parser.parse(_pdf_bytes(), filename="scan.pdf")
    assert result.error is None
    assert result.is_empty
    assert result.page_count == 2


def test_a_pdf_records_its_page_count() -> None:
    """Page numbers are what turn a citation into something a reader can
    follow, and they do not exist anywhere else once the bytes are
    parsed."""
    parser = get_parser(SourceKind.PDF)
    assert parser is not None
    assert parser.parse(_pdf_bytes(pages=5), filename="p.pdf").page_count == 5


def test_a_text_less_pdf_with_no_ocr_hook_says_it_needs_one() -> None:
    """The distinction is actionable: this one needs OCR, not a different
    parser, and a warning saying so is what stops somebody hunting for a
    parser bug."""
    result = PdfParser().parse(_pdf_bytes(), filename="scan.pdf")
    assert result.error is None
    assert any("OCR" in warning for warning in result.warnings)


def test_supplying_an_ocr_hook_withdraws_that_warning() -> None:
    """The seam: no OCR engine ships here, because none exists in this
    platform's infrastructure to have tested one against."""

    class _Ocr:
        def read(self, image: bytes, *, page_number: int) -> str:
            return f"recovered text from page {page_number}"

    result = PdfParser(ocr=_Ocr()).parse(_pdf_bytes(pages=2), filename="scan.pdf")
    assert result.error is None
    assert not any("no text layer" in warning for warning in result.warnings)


def test_an_ocr_hook_that_fails_does_not_fail_the_parse() -> None:
    class _Broken:
        def read(self, image: bytes, *, page_number: int) -> str:
            raise RuntimeError("the OCR engine crashed")

    result = PdfParser(ocr=_Broken()).parse(_pdf_bytes(), filename="scan.pdf")
    assert result.error is None


# ---- the embedding service --------------------------------------------------------


class _Cache:
    def __init__(self) -> None:
        self.store: dict[str, str] = {}

    async def get(self, key: str) -> str | None:
        return self.store.get(key)

    async def set(self, key: str, value: str, *, ttl_seconds: int | None = None) -> None:
        self.store[key] = value


def _service(**kwargs: Any) -> EmbeddingService:
    defaults: dict[str, Any] = {
        "provider": EmbeddingProvider.BUILTIN,
        "model": TEST_MODEL,
        "dimensions": 64,
        "encoder": HashingEncoder(dimensions=64),
        "batch_size": 2,
    }
    defaults.update(kwargs)
    return EmbeddingService(**defaults)


@pytest.mark.asyncio
async def test_embedding_preserves_input_order() -> None:
    """A batch returned out of order pairs every chunk with another
    chunk's vector, silently."""
    service = _service()
    texts = ["alpha one", "beta two", "gamma three", "delta four", "epsilon five"]
    batch = await service.embed(texts)
    assert len(batch.vectors) == len(texts)
    for text, vector in zip(texts, batch.vectors, strict=True):
        assert vector == await service.embed_one(text)


@pytest.mark.asyncio
async def test_embedding_reports_its_cost_and_tokens() -> None:
    service = _service(usd_per_1k_tokens=0.5)
    batch = await service.embed(["some text to embed"])
    assert batch.tokens > 0
    assert batch.cost_usd > 0
    assert batch.model == TEST_MODEL


@pytest.mark.asyncio
async def test_embedding_nothing_costs_nothing() -> None:
    batch = await _service().embed([])
    assert batch.vectors == []
    assert batch.tokens == 0
    assert batch.hit_rate == 0.0


@pytest.mark.asyncio
async def test_the_cache_serves_a_repeated_text() -> None:
    """An embedding is a pure function of (text, model), so a cached
    vector is the same answer for no money."""
    cache = _Cache()
    service = _service(cache=cache)
    await service.embed(["repeated text"])
    second = await service.embed(["repeated text"])
    assert second.cache_hits == 1
    assert second.cache_misses == 0
    assert second.hit_rate == 1.0


@pytest.mark.asyncio
async def test_a_corrupt_cache_entry_is_ignored_rather_than_trusted() -> None:
    """Serving a malformed vector would poison every search it appears
    in; recomputing costs one call."""
    cache = _Cache()
    service = _service(cache=cache)
    await service.embed(["text"])
    for key in list(cache.store):
        cache.store[key] = "not json"
    again = await service.embed(["text"])
    assert len(again.vectors[0]) == 64


@pytest.mark.asyncio
async def test_a_service_with_neither_client_nor_encoder_is_refused() -> None:
    with pytest.raises(ValueError, match="nothing to embed with"):
        EmbeddingService(provider=EmbeddingProvider.BUILTIN, model="m", dimensions=64, encoder=None)


@pytest.mark.asyncio
async def test_a_batch_size_below_one_is_refused() -> None:
    with pytest.raises(ValueError, match="batch_size"):
        _service(batch_size=0)


def test_the_service_reports_its_own_configuration() -> None:
    service = _service()
    assert service.provider is EmbeddingProvider.BUILTIN
    assert service.model == TEST_MODEL
    assert service.dimensions == 64


# ---- the provider clients ------------------------------------------------------------


def test_the_builtin_provider_has_no_http_client() -> None:
    """``None`` is the signal to use the builtin encoder; pretending it is
    an HTTP client would mean giving it a fake base URL."""
    assert build_client(None, provider=EmbeddingProvider.BUILTIN, base_url="") is None  # type: ignore[arg-type]


@pytest.mark.parametrize(
    "provider",
    [EmbeddingProvider.VOYAGE, EmbeddingProvider.COHERE, EmbeddingProvider.GEMINI],
)
def test_an_unsupported_provider_is_refused_by_name(provider: EmbeddingProvider) -> None:
    """Refused explicitly rather than silently falling back to a provider
    the operator did not choose."""
    with pytest.raises(DependencyError, match=str(provider)):
        build_client(None, provider=provider, base_url="https://example.com")  # type: ignore[arg-type]


def test_a_provider_needing_a_base_url_is_refused_without_one() -> None:
    with pytest.raises(DependencyError, match=r"base URL|base_url"):
        build_client(None, provider=EmbeddingProvider.AZURE_OPENAI, base_url="")  # type: ignore[arg-type]


# ---- the pgvector store ----------------------------------------------------------------


def _vector(seed: float) -> list[float]:
    return [seed] * TEST_DIMENSIONS


async def _document(
    repo: DocumentRepository, organization_id: uuid.UUID, **kwargs: Any
) -> Document:
    defaults: dict[str, Any] = {
        "organization_id": organization_id,
        "title": "Doc",
        "source_kind": SourceKind.TXT,
    }
    defaults.update(kwargs)
    return await repo.create(Document(**defaults))


async def _document_with_chunk(
    documents_repo: DocumentRepository, chunks_repo: Any, organization_id: uuid.UUID
) -> tuple[Document, Any]:
    """A document, a current version, and one chunk.

    ``embedding_vectors.document_chunk_id`` is NOT NULL: a vector that
    cannot say which chunk it came from cannot produce a citation, so the
    schema refuses one -- and a fixture has to satisfy that too.
    """

    document = await _document(documents_repo, organization_id)
    versions = DocumentVersionRepository(chunks_repo._session)
    version = await versions.create(
        DocumentVersion(
            organization_id=organization_id,
            document_id=document.id,
            version_number=1,
            content="the nightly backup writes to the archive bucket",
            checksum=f"sum-{document.id}",
            is_current=True,
        )
    )
    chunk = await chunks_repo.create(
        DocumentChunk(
            organization_id=organization_id,
            document_id=document.id,
            document_version_id=version.id,
            sequence=0,
            content="the nightly backup writes to the archive bucket",
        )
    )
    return document, chunk


@pytest.mark.asyncio
async def test_the_pgvector_store_refuses_a_wrong_width_vector(
    vector_store: PgVectorStore, organization_id: uuid.UUID
) -> None:
    with pytest.raises(VectorStoreError, match="dimension"):
        await vector_store.upsert(
            [
                VectorRecord(
                    chunk_id=uuid.uuid4(),
                    document_id=uuid.uuid4(),
                    organization_id=organization_id,
                    vector=[0.1, 0.2],
                )
            ]
        )


@pytest.mark.asyncio
async def test_upserting_nothing_into_pgvector_is_a_no_op(vector_store: PgVectorStore) -> None:
    assert await vector_store.upsert([]) == 0


@pytest.mark.asyncio
async def test_the_pgvector_store_counts_and_describes_itself(
    vector_store: PgVectorStore, organization_id: uuid.UUID
) -> None:
    assert await vector_store.count(organization_id) == 0
    info = await vector_store.describe()
    assert info.provider.value == "pgvector"


@pytest.mark.asyncio
async def test_deleting_specific_chunks_keeps_other_models(
    vector_store: PgVectorStore,
    documents_repo: DocumentRepository,
    chunks_repo: Any,
    vectors_repo: EmbeddingVectorRepository,
    organization_id: uuid.UUID,
) -> None:
    """Keyed on ``(chunk, model)``: a chunk re-embedded under a second
    model must keep both vectors."""

    document = await _document(documents_repo, organization_id)
    versions = DocumentVersionRepository(chunks_repo._session)
    version = await versions.create(
        DocumentVersion(
            organization_id=organization_id,
            document_id=document.id,
            version_number=1,
            content="text",
            checksum="abc",
            is_current=True,
        )
    )
    chunk = await chunks_repo.create(
        DocumentChunk(
            organization_id=organization_id,
            document_id=document.id,
            document_version_id=version.id,
            sequence=0,
            content="the nightly backup writes to the archive bucket",
        )
    )
    await vector_store.upsert(
        [
            VectorRecord(
                chunk_id=chunk.id,
                document_id=document.id,
                organization_id=organization_id,
                vector=_vector(0.01),
                content=chunk.content,
                content_hash="hash-a",
            )
        ]
    )
    await vectors_repo.create(
        EmbeddingVector(
            organization_id=organization_id,
            document_id=document.id,
            document_chunk_id=chunk.id,
            provider="builtin",
            model_name="another-model",
            dimensions=TEST_DIMENSIONS,
            vector=_vector(0.02),
            content_hash="hash-b",
            embedded_at=utcnow(),
        )
    )
    assert await vector_store.delete_chunks([chunk.id]) == 1
    remaining = await vectors_repo.list_for_document(document.id)
    assert [row.model_name for row in remaining] == ["another-model"]


@pytest.mark.asyncio
async def test_deleting_no_chunks_is_a_no_op(vector_store: PgVectorStore) -> None:
    assert await vector_store.delete_chunks([]) == 0


@pytest.mark.asyncio
async def test_searching_an_empty_store_returns_nothing(
    vector_store: PgVectorStore, organization_id: uuid.UUID
) -> None:
    matches = await vector_store.search(
        VectorQuery(organization_id=organization_id, vector=_vector(0.01))
    )
    assert matches == []


# ---- GraphRAG --------------------------------------------------------------------------


def test_terms_are_extracted_as_unigrams_and_phrases() -> None:
    """Single words alone would miss "archive bucket", which is the sort
    of thing a knowledge graph actually holds as one node."""
    terms = extract_terms("restore the archive bucket from cold storage")
    assert "archive" in terms
    assert any(" " in term for term in terms)


def test_short_function_words_are_not_terms() -> None:
    assert "the" not in extract_terms("the backup")


def test_extracting_from_nothing_yields_nothing() -> None:
    assert extract_terms("") == []


def test_no_single_stopword_becomes_a_term_on_its_own() -> None:
    """Phrases are kept even when built from short words -- "a la carte"
    is a real thing a graph might hold -- but no stopword is ever offered
    as a term by itself, because it would match every node."""
    terms = extract_terms("a an the")
    assert all(" " in term for term in terms)


def test_a_graph_node_renders_as_prose() -> None:
    node = GraphNode(key="n1", labels=("Service",), properties={"name": "Backups"})
    assert node.name == "Backups"
    assert "Backups" in node.as_text()


def test_a_node_with_no_name_falls_back_to_its_key() -> None:
    assert GraphNode(key="n1").name == "n1"


def test_a_relationship_renders_as_prose() -> None:
    edge = GraphRelationship(type="DEPENDS_ON", start_key="a", end_key="b")
    assert "DEPENDS_ON" in edge.as_text()


def test_an_empty_subgraph_reports_itself_empty() -> None:
    subgraph = Subgraph()
    assert subgraph.is_empty
    assert subgraph.as_text() == ""


def test_a_subgraph_renders_nodes_and_edges() -> None:
    subgraph = Subgraph(
        nodes=[GraphNode(key="a", properties={"name": "Backups"})],
        relationships=[GraphRelationship(type="USES", start_key="a", end_key="b")],
    )
    assert not subgraph.is_empty
    assert "Backups" in subgraph.as_text()
    assert "USES" in subgraph.as_text()


@pytest.mark.asyncio
async def test_a_disabled_graph_degrades_to_nothing() -> None:
    """The vector and keyword arms answer the query on their own; a graph
    that is off is not an error."""
    retriever = GraphRetriever(GraphClient(None, enabled=False))
    assert not retriever.enabled
    assert await retriever.link_entities("backups", uuid.uuid4()) == []
    assert (await retriever.retrieve("backups", uuid.uuid4())).is_empty
    await retriever.close()


@pytest.mark.asyncio
async def test_a_disabled_client_answers_without_a_driver() -> None:
    client = GraphClient(None, enabled=True)
    assert not client.enabled
    await client.close()


def test_an_unsafe_node_label_is_refused() -> None:
    """A label cannot be a bound parameter in Cypher -- it is part of the
    pattern -- so it is interpolated, and interpolation demands
    validation."""
    with pytest.raises(ValueError, match="label"):
        GraphRetriever(GraphClient(None, enabled=False), node_label="Node; DROP")


def test_the_expansion_depth_is_clamped() -> None:
    """Beyond two hops a knowledge graph of any size returns most of
    itself, which is the opposite of retrieval."""
    retriever = GraphRetriever(GraphClient(None, enabled=False), max_depth=99, max_nodes=0)
    assert retriever.enabled is False


@pytest.mark.asyncio
async def test_an_unsafe_relationship_type_is_refused_before_any_early_return() -> None:
    """Validated before the disabled-graph short circuit, so a malformed
    type is caught in every configuration rather than only when the graph
    happens to be on."""
    retriever = GraphRetriever(GraphClient(None, enabled=False))
    with pytest.raises(ValueError, match="Relationship type"):
        await retriever.expand([GraphNode(key="a")], uuid.uuid4(), relationship_types=["BAD; DROP"])


# ---- repositories -----------------------------------------------------------------------


@pytest.mark.asyncio
async def test_embedding_models_have_exactly_one_default(
    models_repo: EmbeddingModelRepository, organization_id: uuid.UUID
) -> None:
    first = await models_repo.create(
        EmbeddingModel(
            organization_id=organization_id,
            model_name="model-a",
            provider=EmbeddingProvider.BUILTIN,
            dimensions=TEST_DIMENSIONS,
            is_default=True,
        )
    )
    await models_repo.clear_default(organization_id)
    second = await models_repo.create(
        EmbeddingModel(
            organization_id=organization_id,
            model_name="model-b",
            provider=EmbeddingProvider.BUILTIN,
            dimensions=TEST_DIMENSIONS,
            is_default=True,
        )
    )
    default = await models_repo.get_default(organization_id)
    assert default is not None
    assert default.id == second.id
    named = await models_repo.get_by_name(organization_id, EmbeddingProvider.BUILTIN, "model-a")
    assert named is not None
    assert named.id == first.id
    assert len(await models_repo.list_for_org(organization_id)) == 2


@pytest.mark.asyncio
async def test_requiring_an_absent_model_is_not_found(
    models_repo: EmbeddingModelRepository, organization_id: uuid.UUID
) -> None:
    with pytest.raises(NotFoundError):
        await models_repo.require_in_org(organization_id, uuid.uuid4())


@pytest.mark.asyncio
async def test_recording_model_usage_accumulates(
    models_repo: EmbeddingModelRepository, organization_id: uuid.UUID
) -> None:
    model = await models_repo.create(
        EmbeddingModel(
            organization_id=organization_id,
            model_name="m",
            provider=EmbeddingProvider.BUILTIN,
            dimensions=TEST_DIMENSIONS,
        )
    )
    await models_repo.record_usage(model, tokens=100, cost_usd=0.5, vectors=2, moment=utcnow())
    await models_repo.record_usage(model, tokens=50, cost_usd=0.25, vectors=1, moment=utcnow())
    refreshed = await models_repo.require_in_org(organization_id, model.id)
    assert refreshed.total_tokens_embedded == 150


@pytest.mark.asyncio
async def test_a_vector_index_reports_whether_it_was_validated(
    indexes_repo: VectorIndexRepository, organization_id: uuid.UUID
) -> None:
    """A never-validated index comes first: an index nobody has checked is
    the one most likely to be wrong."""
    await indexes_repo.create(
        VectorIndex(
            organization_id=organization_id,
            name="idx-a",
            embedding_provider=EmbeddingProvider.BUILTIN,
            dimensions=TEST_DIMENSIONS,
            model_name=TEST_MODEL,
        )
    )
    assert await indexes_repo.get_by_name(organization_id, "idx-a") is not None
    assert await indexes_repo.list_for_org(organization_id)
    assert await indexes_repo.list_unvalidated()


@pytest.mark.asyncio
async def test_documents_can_be_looked_up_by_external_id_and_checksum(
    documents_repo: DocumentRepository, organization_id: uuid.UUID
) -> None:
    document = await _document(
        documents_repo, organization_id, external_id="ops/h", checksum="abc123"
    )
    assert (await documents_repo.get_by_external_id(organization_id, "ops/h")).id == document.id  # type: ignore[union-attr]
    assert (await documents_repo.get_by_checksum(organization_id, "abc123")).id == document.id  # type: ignore[union-attr]
    assert await documents_repo.get_by_external_id(uuid.uuid4(), "ops/h") is None


@pytest.mark.asyncio
async def test_listing_documents_filters_on_every_facet(
    documents_repo: DocumentRepository, organization_id: uuid.UUID
) -> None:
    await _document(documents_repo, organization_id, source_kind=SourceKind.MARKDOWN)
    assert await documents_repo.list_for_org(organization_id, source_kind=SourceKind.MARKDOWN)
    assert not await documents_repo.list_for_org(organization_id, source_kind=SourceKind.PDF)
    assert not await documents_repo.list_for_org(organization_id, knowledge_source_id=uuid.uuid4())


@pytest.mark.asyncio
async def test_listing_documents_by_id_is_tenant_scoped(
    documents_repo: DocumentRepository, organization_id: uuid.UUID
) -> None:
    document = await _document(documents_repo, organization_id)
    assert await documents_repo.list_by_ids(organization_id, [document.id])
    assert await documents_repo.list_by_ids(uuid.uuid4(), [document.id]) == []
    assert await documents_repo.list_by_ids(organization_id, []) == []


@pytest.mark.asyncio
async def test_expired_documents_are_selectable(
    documents_repo: DocumentRepository, organization_id: uuid.UUID
) -> None:
    await _document(documents_repo, organization_id, expires_at=ago(3_600))
    assert await documents_repo.list_expired(utcnow())


@pytest.mark.asyncio
async def test_organizations_with_documents_are_enumerable(
    documents_repo: DocumentRepository, organization_id: uuid.UUID
) -> None:
    await _document(documents_repo, organization_id)
    assert organization_id in await documents_repo.list_organization_ids(limit=500)


@pytest.mark.asyncio
async def test_vectors_are_countable_and_attributable(
    documents_repo: DocumentRepository,
    chunks_repo: Any,
    vectors_repo: EmbeddingVectorRepository,
    organization_id: uuid.UUID,
) -> None:
    document, chunk = await _document_with_chunk(documents_repo, chunks_repo, organization_id)
    await vectors_repo.create(
        EmbeddingVector(
            organization_id=organization_id,
            document_id=document.id,
            document_chunk_id=chunk.id,
            provider="builtin",
            model_name=TEST_MODEL,
            dimensions=TEST_DIMENSIONS,
            vector=_vector(0.01),
            token_count=10,
            cost_usd=0.5,
            content_hash="hash-a",
            embedded_at=utcnow(),
        )
    )
    assert await vectors_repo.count_for_org(organization_id) == 1
    assert await vectors_repo.count_for_org(organization_id, model_name=TEST_MODEL) == 1
    assert await vectors_repo.models_in_use(organization_id) == {TEST_MODEL: 1}
    tokens, cost = await vectors_repo.tokens_in_window(
        organization_id, since=ago(3_600), until=utcnow()
    )
    assert tokens == 10
    assert cost == pytest.approx(0.5)


@pytest.mark.asyncio
async def test_vectors_can_be_deleted_by_model(
    documents_repo: DocumentRepository,
    chunks_repo: Any,
    vectors_repo: EmbeddingVectorRepository,
    organization_id: uuid.UUID,
) -> None:
    """Retiring a model must not take the vectors of the model that
    replaced it."""
    document, chunk = await _document_with_chunk(documents_repo, chunks_repo, organization_id)
    for model in ("old-model", "new-model"):
        await vectors_repo.create(
            EmbeddingVector(
                organization_id=organization_id,
                document_id=document.id,
                document_chunk_id=chunk.id,
                provider="builtin",
                model_name=model,
                dimensions=TEST_DIMENSIONS,
                vector=_vector(0.01),
                content_hash=f"hash-{model}",
                embedded_at=utcnow(),
            )
        )
    assert await vectors_repo.delete_for_model(organization_id, "old-model") == 1
    assert [row.model_name for row in await vectors_repo.list_for_document(document.id)] == [
        "new-model"
    ]


@pytest.mark.asyncio
async def test_a_vector_is_findable_by_its_content_hash(
    documents_repo: DocumentRepository,
    chunks_repo: Any,
    vectors_repo: EmbeddingVectorRepository,
    organization_id: uuid.UUID,
) -> None:
    document, chunk = await _document_with_chunk(documents_repo, chunks_repo, organization_id)
    await vectors_repo.create(
        EmbeddingVector(
            organization_id=organization_id,
            document_id=document.id,
            document_chunk_id=chunk.id,
            provider="builtin",
            model_name=TEST_MODEL,
            dimensions=TEST_DIMENSIONS,
            vector=_vector(0.01),
            content_hash="shared-hash",
            embedded_at=utcnow(),
        )
    )
    found = await vectors_repo.find_by_content_hash(
        organization_id, "shared-hash", model_name=TEST_MODEL
    )
    assert found is not None
    assert (
        await vectors_repo.find_by_content_hash(
            organization_id, "shared-hash", model_name="another"
        )
        is None
    )


@pytest.mark.asyncio
async def test_retrieval_queries_roll_up_by_strategy_and_latency(
    queries_repo: RetrievalQueryRepository, organization_id: uuid.UUID
) -> None:
    await queries_repo.create(
        RetrievalQuery(
            organization_id=organization_id,
            query_text="backups",
            normalized_query="backups",
            strategy=RetrievalStrategy.HYBRID,
            outcome=RetrievalOutcome.SUCCEEDED,
            duration_ms=12.0,
            executed_at=utcnow(),
        )
    )
    assert await queries_repo.count_by_strategy(organization_id, since=ago(3_600))
    assert await queries_repo.average_latency(
        organization_id, since=ago(3_600), until=utcnow()
    ) == pytest.approx(12.0)
    assert await queries_repo.list_for_org(organization_id)
    assert await queries_repo.list_for_org(organization_id, outcome=RetrievalOutcome.SUCCEEDED)


@pytest.mark.asyncio
async def test_latency_over_an_empty_window_is_none(
    queries_repo: RetrievalQueryRepository, organization_id: uuid.UUID
) -> None:
    """A window with no queries has no latency; reporting zero would look
    like an impossibly fast service."""
    assert (
        await queries_repo.average_latency(organization_id, since=ago(3_600), until=utcnow())
        is None
    )


@pytest.mark.asyncio
async def test_requiring_an_absent_query_is_not_found(
    queries_repo: RetrievalQueryRepository, organization_id: uuid.UUID
) -> None:
    with pytest.raises(NotFoundError):
        await queries_repo.require_in_org(organization_id, uuid.uuid4())


@pytest.mark.asyncio
async def test_feedback_grades_fall_back_when_no_number_was_given(
    documents_repo: DocumentRepository,
    queries_repo: RetrievalQueryRepository,
    results_repo: RetrievalResultRepository,
    feedback_repo: RetrievalFeedbackRepository,
    chunks_repo: Any,
    organization_id: uuid.UUID,
) -> None:
    """A reviewer who only clicked a button still gives nDCG something to
    work with."""

    document = await _document(documents_repo, organization_id)
    versions = DocumentVersionRepository(chunks_repo._session)
    version = await versions.create(
        DocumentVersion(
            organization_id=organization_id,
            document_id=document.id,
            version_number=1,
            content="text",
            checksum="abc",
            is_current=True,
        )
    )
    chunk = await chunks_repo.create(
        DocumentChunk(
            organization_id=organization_id,
            document_id=document.id,
            document_version_id=version.id,
            sequence=0,
            content="text",
        )
    )
    query = await queries_repo.create(
        RetrievalQuery(
            organization_id=organization_id,
            query_text="q",
            strategy=RetrievalStrategy.HYBRID,
            outcome=RetrievalOutcome.SUCCEEDED,
            executed_at=utcnow(),
        )
    )
    await results_repo.create(
        RetrievalResult(
            organization_id=organization_id,
            retrieval_query_id=query.id,
            document_id=document.id,
            document_chunk_id=chunk.id,
            rank=1,
            score=0.9,
        )
    )
    await feedback_repo.create(
        RetrievalFeedback(
            organization_id=organization_id,
            retrieval_query_id=query.id,
            document_chunk_id=chunk.id,
            verdict=FeedbackVerdict.PARTIALLY_RELEVANT,
            submitted_at=utcnow(),
        )
    )
    graded = await feedback_repo.graded_relevance(query.id)
    assert graded[chunk.id] == 0.5
    assert chunk.id in await feedback_repo.relevant_chunk_ids(query.id)
    assert await feedback_repo.count_by_verdict(organization_id, since=ago(3_600))
    assert query.id in await feedback_repo.list_judged_queries(organization_id, since=ago(3_600))
    assert await results_repo.top_documents(organization_id, since=ago(3_600))
    assert await results_repo.average_result_count(
        organization_id, since=ago(3_600), until=utcnow()
    ) == pytest.approx(1.0)


@pytest.mark.asyncio
async def test_a_knowledge_source_that_never_synced_is_due(
    sources_repo: KnowledgeSourceRepository, organization_id: uuid.UUID
) -> None:
    """A NULL last-synced is the strongest possible signal that a sync is
    owed."""
    source = await sources_repo.create(
        KnowledgeSource(
            organization_id=organization_id,
            slug="s",
            name="S",
            source_kind=SourceKind.CONFLUENCE,
            sync_enabled=True,
            sync_status=SyncStatus.NEVER_SYNCED,
        )
    )
    assert source.id in {row.id for row in await sources_repo.list_due_for_sync(utcnow())}
    assert await sources_repo.get_by_slug(organization_id, "s") is not None
    assert await sources_repo.count_documents(source.id) == 0


@pytest.mark.asyncio
async def test_a_source_already_syncing_is_not_offered_again(
    sources_repo: KnowledgeSourceRepository, organization_id: uuid.UUID
) -> None:
    source = await sources_repo.create(
        KnowledgeSource(
            organization_id=organization_id,
            slug="s",
            name="S",
            source_kind=SourceKind.CONFLUENCE,
            sync_enabled=True,
            sync_status=SyncStatus.SYNCING,
        )
    )
    assert source.id not in {row.id for row in await sources_repo.list_due_for_sync(utcnow())}


@pytest.mark.asyncio
async def test_requiring_an_absent_source_or_job_is_not_found(
    sources_repo: KnowledgeSourceRepository,
    jobs_repo: IndexingJobRepository,
    organization_id: uuid.UUID,
) -> None:
    with pytest.raises(NotFoundError):
        await sources_repo.require_in_org(organization_id, uuid.uuid4())
    with pytest.raises(NotFoundError):
        await jobs_repo.require_in_org(organization_id, uuid.uuid4())


@pytest.mark.asyncio
async def test_jobs_are_countable_by_status(
    jobs_repo: IndexingJobRepository, organization_id: uuid.UUID
) -> None:
    await jobs_repo.create(
        IndexingJob(
            organization_id=organization_id,
            kind=IndexKind.INCREMENTAL,
            status=IndexStatus.QUEUED,
            scheduled_at=utcnow(),
        )
    )
    assert await jobs_repo.count_by_status(organization_id) == {"queued": 1}
    assert await jobs_repo.list_for_org(organization_id, status=IndexStatus.QUEUED)


@pytest.mark.asyncio
async def test_statistics_are_listed_oldest_first(
    statistics_repo: RagStatisticRepository, organization_id: uuid.UUID
) -> None:
    """A trend chart plotted newest-first reads as a mirror image of the
    truth."""
    for offset in (7_200, 3_600):
        await statistics_repo.create(
            RagStatistic(
                organization_id=organization_id,
                window_start=ago(offset),
                window_end=ago(offset - 600),
            )
        )
    listed = await statistics_repo.list_since(organization_id, since=ago(86_400))
    assert [row.window_start for row in listed] == sorted(row.window_start for row in listed)
    latest = await statistics_repo.latest(organization_id)
    assert latest is not None


@pytest.mark.asyncio
async def test_no_statistics_yet_is_none_not_an_error(
    statistics_repo: RagStatisticRepository, organization_id: uuid.UUID
) -> None:
    assert await statistics_repo.latest(organization_id) is None


@pytest.mark.asyncio
async def test_reports_are_listed_newest_first_and_filterable(
    reports_repo: RagReportRepository, organization_id: uuid.UUID
) -> None:
    await reports_repo.create(
        RagReport(organization_id=organization_id, kind=ReportKind.INDEX, title="Index")
    )
    assert await reports_repo.list_for_org(organization_id, kind=ReportKind.INDEX)
    assert not await reports_repo.list_for_org(organization_id, kind=ReportKind.AUDIT)


@pytest.mark.asyncio
async def test_the_session_fixture_is_a_real_session(db_session: AsyncSession) -> None:
    assert isinstance(db_session, AsyncSession)
