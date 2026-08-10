"""PDF and DOCX parsing, plus the OCR extension point.

**Page numbers are the reason PDF gets its own parser** rather than being
funnelled through a generic text extractor. A citation into a 200-page
manual that names only the document is not a reference somebody can
follow, and the page number exists only during extraction -- once the
text is concatenated it is gone for good.

**OCR is an injected hook, not a dependency.** Optical character
recognition needs a Tesseract binary (or a cloud API credential) that
this container does not ship and cannot usefully fake. Rather than
declare a dependency that would fail at runtime on every deployment that
did not separately install it, :class:`OcrHook` is a protocol a
deployment supplies. The same pattern user-management-service uses for
its virus-scanning hook.

Without a hook, a scanned PDF parses **successfully and empty**, and that
is precisely the signal a caller needs: not an error, but "this document
has no text layer, so it needs OCR". Reporting it as a parse failure
would send someone looking for a broken parser.
"""

from __future__ import annotations

import io
from collections.abc import Sequence
from typing import Protocol

import docx
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.models.enums import SourceKind
from app.parsers.base import (
    ParsedBlock,
    ParseResult,
    blocks_to_text,
    oversized,
    register,
)


class OcrHook(Protocol):
    """Optical character recognition, supplied by the deployment.

    Receives the raw bytes of one page image and returns whatever text it
    can read. Implementations are expected to be slow and are called only
    for pages that produced no text layer.
    """

    def read(self, image: bytes, *, page_number: int) -> str:
        """Text recognised in *image*, or an empty string."""
        ...


class PdfParser:
    """PDF, one block per page, page numbers preserved.

    Encrypted PDFs are attempted with the empty-string password, which
    opens the very common case of a document encrypted only to set
    permissions rather than to restrict reading. A genuinely
    password-protected document fails with a clear reason instead of an
    opaque extraction error.
    """

    name = "pdf"

    def __init__(self, ocr: OcrHook | None = None) -> None:
        self._ocr = ocr

    def parse(self, data: bytes, *, filename: str | None = None) -> ParseResult:
        too_big = oversized(data)
        if too_big is not None:
            return too_big

        try:
            reader = PdfReader(io.BytesIO(data))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    return ParseResult(
                        error=(
                            "PDF is password-protected. Supply a decrypted copy; this "
                            "service does not attempt to break document encryption."
                        ),
                        parser=self.name,
                    )
            pages = list(reader.pages)
        except (PdfReadError, OSError, ValueError) as exc:
            return ParseResult(error=f"PDF could not be read: {exc}", parser=self.name)

        blocks: list[ParsedBlock] = []
        warnings: list[str] = []
        ocr_used = 0
        for number, page in enumerate(pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as exc:
                # One unreadable page must not lose the other 199.
                warnings.append(f"Page {number} could not be extracted: {exc}")
                continue
            if not text.strip() and self._ocr is not None:
                text = self._recognise(page, number=number, warnings=warnings)
                if text.strip():
                    ocr_used += 1
            if text.strip():
                blocks.append(ParsedBlock(text=text.strip(), page_number=number))

        metadata = self._metadata(reader)
        if ocr_used:
            warnings.append(f"{ocr_used} page(s) were read by OCR rather than a text layer.")
        if not blocks and self._ocr is None:
            warnings.append(
                "No text layer was found on any page. This is most likely a scanned "
                "document, which needs an OCR hook rather than a different parser."
            )

        return ParseResult(
            text=blocks_to_text(blocks),
            blocks=blocks,
            metadata=metadata,
            page_count=len(pages),
            parser=self.name,
            warnings=warnings,
        )

    def _recognise(self, page: object, *, number: int, warnings: list[str]) -> str:
        """Run the OCR hook over one page's embedded images."""
        if self._ocr is None:  # pragma: no cover - guarded by the caller
            return ""
        recognised: list[str] = []
        try:
            images = list(getattr(page, "images", []))
        except Exception as exc:
            warnings.append(f"Page {number} images could not be read: {exc}")
            return ""
        for image in images:
            payload = getattr(image, "data", None)
            if not payload:
                continue
            try:
                recognised.append(self._ocr.read(payload, page_number=number))
            except Exception as exc:
                warnings.append(f"OCR failed on page {number}: {exc}")
        return "\n".join(part for part in recognised if part.strip())

    @staticmethod
    def _metadata(reader: object) -> dict[str, str]:
        """Title, author, and subject where the document declares them."""
        raw = getattr(reader, "metadata", None)
        if not raw:
            return {}
        fields = {"title": "/Title", "author": "/Author", "subject": "/Subject"}
        found: dict[str, str] = {}
        for name, key in fields.items():
            try:
                value = raw.get(key)
            except Exception:  # pragma: no cover - defensive against odd metadata
                continue
            if value:
                found[name] = str(value).strip()
        return found


class DocxParser:
    """DOCX, with heading trails and tables preserved.

    Word's own heading styles carry the document's outline, so the
    section trail is recovered from ``Heading 1``..``Heading 6`` rather
    than guessed from formatting. Tables are emitted as labelled rows for
    the same reason CSV is -- a bare row of cells does not say what its
    columns mean.
    """

    name = "docx"

    def parse(self, data: bytes, *, filename: str | None = None) -> ParseResult:
        too_big = oversized(data)
        if too_big is not None:
            return too_big

        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:
            return ParseResult(error=f"DOCX could not be read: {exc}", parser=self.name)

        blocks: list[ParsedBlock] = []
        trail: list[str] = []
        for paragraph in document.paragraphs:
            content = paragraph.text.strip()
            if not content:
                continue
            level = _heading_level(paragraph.style.name if paragraph.style else None)
            if level:
                del trail[level - 1 :]
                trail.append(content)
                continue
            blocks.append(
                ParsedBlock(
                    text=content,
                    section_path=tuple(trail),
                    heading=trail[-1] if trail else None,
                )
            )

        blocks.extend(_table_blocks(document.tables, trail=tuple(trail)))

        return ParseResult(
            text=blocks_to_text(blocks),
            blocks=blocks,
            metadata=_docx_metadata(document),
            parser=self.name,
        )


MAX_HEADING_LEVEL = 6
"""Word offers Heading 1 through Heading 6; anything deeper is a custom
style whose outline position this parser cannot know."""


def _heading_level(style_name: str | None) -> int | None:
    """The outline level of a Word style, or ``None`` if it is not one."""
    if not style_name:
        return None
    lowered = style_name.strip().lower()
    if lowered == "title":
        return 1
    if lowered.startswith("heading "):
        suffix = lowered.removeprefix("heading ").strip()
        if suffix.isdigit():
            level = int(suffix)
            return level if 1 <= level <= MAX_HEADING_LEVEL else None
    return None


def _table_blocks(tables: Sequence[object], *, trail: tuple[str, ...]) -> list[ParsedBlock]:
    """Flatten Word tables into labelled rows."""
    blocks: list[ParsedBlock] = []
    for table in tables:
        rows = list(getattr(table, "rows", []))
        if not rows:
            continue
        header = [cell.text.strip() for cell in rows[0].cells]
        for row in rows[1:]:
            pairs = [
                f"{header[index] or f'column {index + 1}'}: {cell.text.strip()}"
                for index, cell in enumerate(row.cells)
                if index < len(header) and cell.text.strip()
            ]
            if pairs:
                blocks.append(ParsedBlock(text=", ".join(pairs), section_path=trail, is_table=True))
    return blocks


def _docx_metadata(document: object) -> dict[str, str]:
    """Core properties Word records on the document."""
    properties = getattr(document, "core_properties", None)
    if properties is None:  # pragma: no cover - always present in python-docx
        return {}
    found: dict[str, str] = {}
    for name in ("title", "author", "subject", "category"):
        value = getattr(properties, name, None)
        if value:
            found[name] = str(value).strip()
    return found


register(SourceKind.PDF, PdfParser)
register(SourceKind.DOCX, DocxParser)


__all__ = ["DocxParser", "OcrHook", "PdfParser"]
