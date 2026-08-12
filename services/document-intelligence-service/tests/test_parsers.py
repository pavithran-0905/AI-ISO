"""Tests for format detection and every parser.

Real bytes throughout: a PDF built by pypdf, a DOCX built by python-docx,
an XLSX built by openpyxl, images built by Pillow, and a real ZIP. A parser
test against a hand-written string is a test of the string.
"""

from __future__ import annotations

import io
import json
import zipfile

import pytest

from app.documents.detection import MAX_SNIFF_BYTES, detect_format
from app.documents.parser import (
    DocumentParseError,
    UnsupportedFormatError,
    merge,
    paginate,
    parse,
    parser_for,
    register,
    supported_formats,
)
from app.documents.text_formats import decode, flatten
from app.models.enums import DocumentFormat

TEXT = b"Change request CHG-004821.\nApproved by R. Mehta on 2026-03-14.\n"
MARKDOWN = b"# Runbook\n\n## Rollback\n\n- Redeploy 4.2.0\n- Verify p99 latency\n"
HTML = (
    b"<!doctype html><html><head><title>Postmortem</title>"
    b"<meta name='author' content='R. Mehta'></head><body>"
    b"<script>bad()</script><h1>Outage</h1><p>The API was down 41 minutes.</p>"
    b"<table><tr><td>a</td></tr></table></body></html>"
)
RTF = rb"{\rtf1\ansi{\fonttbl{\f0 Arial;}}\f0\fs24 Change \b request\b0\par Approved: yes\par}"
CSV = b"change_id,system,risk\nCHG-004821,payments-api,high\nCHG-004822,billing,medium\n"
TSV = b"change_id\tsystem\trisk\nCHG-004821\tpayments-api\thigh\n"
JSON_DOC = json.dumps(
    {"change": {"id": "CHG-004821", "risk": "high", "approvers": ["R. Mehta", "A. Novak"]}}
).encode()
YAML_DOC = (
    b"change:\n  id: CHG-004821\n  risk: high\n  approvers:\n    - R. Mehta\n" b"---\nsecond: doc\n"
)
XML_DOC = b"<?xml version='1.0'?><change><id>CHG-004821</id><risk>high</risk></change>"


# ---- detection ----------------------------------------------------------------------


@pytest.mark.parametrize(
    ("data", "filename", "content_type", "expected"),
    [
        (TEXT, "note.txt", "text/plain", DocumentFormat.TXT),
        (MARKDOWN, "runbook.md", None, DocumentFormat.MARKDOWN),
        (HTML, "post.html", "text/html", DocumentFormat.HTML),
        (RTF, None, None, DocumentFormat.RTF),
        (CSV, "changes.csv", "text/csv", DocumentFormat.CSV),
        (TSV, "changes.tsv", None, DocumentFormat.CSV),
        (JSON_DOC, None, "text/plain", DocumentFormat.JSON),
        (YAML_DOC, "change.yaml", None, DocumentFormat.YAML),
        (XML_DOC, None, None, DocumentFormat.XML),
    ],
)
def test_detection(
    data: bytes, filename: str | None, content_type: str | None, expected: DocumentFormat
) -> None:
    guess = detect_format(data, filename=filename, content_type=content_type)
    assert guess.format is expected
    assert guess.evidence


def test_the_bytes_outrank_the_extension() -> None:
    """A .txt holding a PDF signature is a PDF that was renamed."""
    guess = detect_format(b"%PDF-1.7\nstuff", filename="notes.txt", content_type="text/plain")
    assert guess.format is DocumentFormat.PDF


def test_content_beats_a_generic_declared_type() -> None:
    """Browsers send text/plain for anything they do not recognise."""
    assert detect_format(JSON_DOC, content_type="text/plain").format is DocumentFormat.JSON


def test_an_unidentifiable_payload_is_unknown_not_a_guess() -> None:
    guess = detect_format(b"\x00\x01\x02\x03\x04\x05\x06\x07", filename="mystery.bin")
    assert guess.format is DocumentFormat.UNKNOWN
    assert guess.is_known is False


def test_an_empty_payload_is_unknown() -> None:
    guess = detect_format(b"")
    assert guess.format is DocumentFormat.UNKNOWN
    assert "empty" in guess.evidence


def test_text_with_no_named_format_falls_back_to_txt() -> None:
    guess = detect_format(b"just some prose with no markers at all")
    assert guess.format is DocumentFormat.TXT
    assert guess.confidence < 0.6


def test_the_sniff_window_is_bounded() -> None:
    """Detection must not read a hundred megabytes to name a format."""
    padded = b"%PDF-1.7\n" + b"x" * (MAX_SNIFF_BYTES * 4)
    assert detect_format(padded).format is DocumentFormat.PDF


# ---- text parsers -------------------------------------------------------------------


def test_txt_parses_and_paginates() -> None:
    document = parse(TEXT, filename="note.txt")
    assert document.format is DocumentFormat.TXT
    assert "CHG-004821" in document.text
    assert document.page_count == 1
    assert document.word_count > 5


def test_markdown_keeps_its_markup() -> None:
    """Headings and bullets are the structure downstream engines look for."""
    document = parse(MARKDOWN, filename="runbook.md")
    assert "# Runbook" in document.text
    assert document.metadata["heading_count"] == "2"


def test_html_drops_scripts_and_keeps_metadata() -> None:
    document = parse(HTML, filename="post.html")
    assert "bad()" not in document.text
    assert "The API was down 41 minutes." in document.text
    assert document.metadata["title"] == "Postmortem"
    assert document.metadata["meta:author"] == "R. Mehta"


def test_rtf_strips_the_font_table() -> None:
    """The pattern counted the keyword's backslash twice and leaked "Arial;"."""
    document = parse(RTF)
    assert "Arial" not in document.text
    assert "Change request" in document.text
    assert document.warnings


def test_csv_sniffs_its_delimiter() -> None:
    comma = parse(CSV, filename="c.csv")
    tab = parse(TSV, filename="c.tsv")
    assert comma.metadata["row_count"] == "3"
    assert tab.metadata["delimiter"] == repr("\t")
    assert "payments-api" in tab.text


def test_a_malformed_delimited_file_degrades_to_text() -> None:
    broken = b'a,b\n"unclosed quote,c\n'
    document = parse(broken, filename="broken.csv")
    assert document.text


def test_json_flattens_to_readable_paths() -> None:
    document = parse(JSON_DOC, filename="c.json")
    assert "change.id: CHG-004821" in document.text
    assert "change.approvers[0]: R. Mehta" in document.text


def test_invalid_json_degrades_to_text_with_a_warning() -> None:
    document = parse(b'{"unclosed": ', filename="broken.json", content_type="application/json")
    assert document.warnings
    assert document.text


def test_yaml_handles_multiple_documents() -> None:
    document = parse(YAML_DOC, filename="c.yaml")
    assert document.metadata["yaml_documents"] == "2"
    assert "change.id: CHG-004821" in document.text


def test_invalid_yaml_degrades_to_text() -> None:
    document = parse(b"key: [unclosed\n", filename="broken.yaml")
    assert document.warnings


def test_xml_renders_element_paths() -> None:
    document = parse(XML_DOC, filename="c.xml")
    assert "change.id: CHG-004821" in document.text
    assert document.metadata["root_element"] == "change"


def test_xml_with_no_elements_degrades_to_text() -> None:
    document = parse(b"not xml at all", fmt=DocumentFormat.XML)
    assert document.warnings


def test_decode_reports_the_encoding_that_worked() -> None:
    text, encoding = decode("café".encode("cp1252"))
    assert "caf" in text
    assert encoding in {"cp1252", "latin-1", "utf-16"}


def test_a_non_utf8_payload_warns_about_its_encoding() -> None:
    document = parse("café".encode("cp1252") + b" change request\n", fmt=DocumentFormat.TXT)
    assert document.metadata["encoding"]


def test_flatten_caps_its_nesting_depth() -> None:
    """A structure nested past the cap is generated or hostile."""
    from app.documents.parser import ParsedDocument

    deep: object = "bottom"
    for _ in range(80):
        deep = {"next": deep}
    document = ParsedDocument(format=DocumentFormat.JSON)
    lines = flatten(deep, document)
    assert lines
    assert any("deeper than" in warning for warning in document.warnings)


# ---- binary parsers -----------------------------------------------------------------


def test_a_pdf_with_no_text_layer_needs_ocr_rather_than_reading_as_empty() -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    writer.add_blank_page(width=595, height=842)
    buffer = io.BytesIO()
    writer.write(buffer)

    document = parse(buffer.getvalue(), filename="scan.pdf")
    assert document.page_count == 2
    assert document.needs_ocr is True
    assert document.pages[0].width == pytest.approx(595, abs=1)
    assert any("text layer" in warning for warning in document.warnings)


def test_an_unreadable_pdf_raises_rather_than_returning_nothing() -> None:
    with pytest.raises(DocumentParseError):
        parse(b"%PDF-1.7\nnot actually a pdf", filename="broken.pdf")


def test_docx_emits_tables_inline_in_body_order() -> None:
    """python-docx exposes paragraphs and tables separately; reading them in
    turn would put every table after every paragraph."""
    import docx

    source = docx.Document()
    source.add_heading("Change Request", level=1)
    source.add_paragraph("Requested by: R. Mehta")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "System"
    table.cell(0, 1).text = "Risk"
    table.cell(1, 0).text = "payments-api"
    table.cell(1, 1).text = "high"
    source.add_paragraph("Rollback tested: yes")
    buffer = io.BytesIO()
    source.save(buffer)

    payload = buffer.getvalue()
    assert detect_format(payload, filename="cr.docx").format is DocumentFormat.DOCX
    document = parse(payload, filename="cr.docx")
    text = document.text
    assert text.index("Requested by") < text.index("| System")
    assert text.index("| System") < text.index("Rollback tested")
    assert document.metadata["table_count"] == "1"


def test_an_unreadable_docx_raises() -> None:
    with pytest.raises(DocumentParseError):
        parse(b"PK\x03\x04garbage-not-a-docx", fmt=DocumentFormat.DOCX)


def test_xlsx_gives_one_page_per_sheet() -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Changes"
    sheet.append(["change_id", "risk"])
    sheet.append(["CHG-004821", "high"])
    workbook.create_sheet("Notes").append(["Rollback tested"])
    buffer = io.BytesIO()
    workbook.save(buffer)

    payload = buffer.getvalue()
    assert detect_format(payload, filename="c.xlsx").format is DocumentFormat.XLSX
    document = parse(payload, filename="c.xlsx")
    assert document.page_count == 2
    assert document.metadata["sheet_count"] == "2"
    assert "CHG-004821" in document.text


def test_an_unreadable_workbook_raises() -> None:
    with pytest.raises(DocumentParseError):
        parse(b"PK\x03\x04not-a-workbook", fmt=DocumentFormat.XLSX)


def test_an_image_needs_ocr_and_reports_its_geometry() -> None:
    from PIL import Image

    buffer = io.BytesIO()
    Image.new("RGB", (640, 480), "white").save(buffer, format="PNG")
    document = parse(buffer.getvalue(), filename="scan.png")
    assert document.needs_ocr is True
    assert document.pages[0].width == 640
    assert document.pages[0].has_text_layer is False
    assert document.metadata["image_format"] == "PNG"


def test_a_multi_frame_tiff_is_one_page_per_frame() -> None:
    """Collapsing a forty-page fax to one page loses every page reference."""
    from PIL import Image

    frames = [Image.new("L", (200, 300), shade) for shade in (255, 200, 128)]
    buffer = io.BytesIO()
    frames[0].save(buffer, format="TIFF", save_all=True, append_images=frames[1:])
    document = parse(buffer.getvalue(), filename="fax.tiff")
    assert document.page_count == 3
    assert document.metadata["frame_count"] == "3"


def test_an_undecodable_image_raises() -> None:
    with pytest.raises(DocumentParseError):
        parse(b"\x89PNG\r\n\x1a\ngarbage", fmt=DocumentFormat.IMAGE)


def _archive(entries: dict[str, bytes], *, compress: bool = True) -> bytes:
    buffer = io.BytesIO()
    mode = zipfile.ZIP_DEFLATED if compress else zipfile.ZIP_STORED
    with zipfile.ZipFile(buffer, "w", mode) as archive:
        for name, payload in entries.items():
            archive.writestr(name, payload)
    return buffer.getvalue()


def test_an_archive_parses_each_member_and_reports_what_it_skipped() -> None:
    payload = _archive(
        {
            "notes.txt": TEXT,
            "changes.csv": CSV,
            "data.json": JSON_DOC,
            "tool.exe": b"MZ" + b"\x00" * 100,
            "mystery.dat": b"\x00\x01\x02\x03\x04\x05\x06\x07\x08",
        }
    )
    document = parse(payload, filename="bundle.zip")
    assert document.metadata["member_count"] == "5"
    assert len(document.attachments) == 3
    assert document.page_count == 3
    assert any("executable" in warning for warning in document.warnings)
    assert any("unrecognised" in warning for warning in document.warnings)


def test_a_nested_archive_is_expanded_one_level() -> None:
    inner = _archive({"deep.txt": b"Deeply nested content here."})
    document = parse(_archive({"inner.zip": inner}), filename="outer.zip")
    assert "Deeply nested" in document.text


def test_a_zip_bomb_is_refused_before_anything_is_read() -> None:
    bomb = _archive({"big.txt": b"A" * (5 * 1024 * 1024)})
    with pytest.raises(DocumentParseError, match="zip bomb"):
        parse(bomb, filename="bomb.zip")


def test_too_many_members_is_refused() -> None:
    payload = _archive({f"f{index}.txt": b"x" for index in range(600)}, compress=False)
    with pytest.raises(DocumentParseError, match="members"):
        parse(payload, filename="many.zip")


def test_an_unreadable_archive_raises() -> None:
    with pytest.raises(DocumentParseError):
        parse(b"PK\x03\x04truncated", fmt=DocumentFormat.ZIP)


# ---- registry and helpers -----------------------------------------------------------


def test_a_second_parser_for_one_format_is_refused() -> None:
    """Whichever module imported last would otherwise win."""

    def other(_data: bytes) -> object:  # pragma: no cover -- never called
        raise AssertionError

    with pytest.raises(ValueError, match="already registered"):
        register(DocumentFormat.TXT, other)  # type: ignore[arg-type]


def test_re_registering_the_same_parser_is_allowed() -> None:
    """Import order can legitimately re-run a module."""
    register(DocumentFormat.TXT, parser_for(DocumentFormat.TXT))


def test_an_unregistered_format_names_what_is_registered() -> None:
    with pytest.raises(UnsupportedFormatError, match="registered formats"):
        parser_for(DocumentFormat.UNKNOWN)


def test_parse_refuses_an_unidentifiable_payload() -> None:
    with pytest.raises(DocumentParseError, match="could not be"):
        parse(b"\x00\x01\x02\x03\x04\x05\x06\x07", filename="x.bin")


def test_parse_records_its_evidence_and_filename() -> None:
    document = parse(TEXT, filename="note.txt")
    assert document.metadata["filename"] == "note.txt"
    assert document.metadata["format_evidence"]


def test_an_explicit_format_skips_detection() -> None:
    document = parse(TEXT, fmt=DocumentFormat.TXT)
    assert "format_evidence" not in document.metadata


def test_paginate_prefers_form_feeds_over_a_character_count() -> None:
    pages = paginate("one\fTwo\fThree", per_page=2)
    assert [page.number for page in pages] == [1, 2, 3]
    assert pages[1].text == "Two"


def test_paginate_splits_on_a_character_count_when_there_are_no_form_feeds() -> None:
    assert len(paginate("abcdefghij", per_page=4)) == 3
    assert len(paginate("short", per_page=None)) == 1


def test_merge_renumbers_pages_continuously() -> None:
    """A citation to "page 4" must mean one thing."""
    first = parse(TEXT, fmt=DocumentFormat.TXT)
    second = parse(MARKDOWN, fmt=DocumentFormat.MARKDOWN)
    merged = merge([first, second], DocumentFormat.TXT)
    assert [page.number for page in merged.pages] == [1, 2]
    assert merged.word_count == first.word_count + second.word_count


def test_merge_propagates_the_ocr_flag() -> None:
    from PIL import Image

    buffer = io.BytesIO()
    Image.new("RGB", (10, 10)).save(buffer, format="PNG")
    scan = parse(buffer.getvalue(), fmt=DocumentFormat.IMAGE)
    text = parse(TEXT, fmt=DocumentFormat.TXT)
    assert merge([text, scan], DocumentFormat.PDF).needs_ocr is True


def test_every_format_but_unknown_is_registered() -> None:
    assert set(supported_formats()) == set(DocumentFormat) - {DocumentFormat.UNKNOWN}
