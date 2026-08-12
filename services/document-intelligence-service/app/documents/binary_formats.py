"""Parsers for the binary formats.

PDF, DOCX, XLSX, images, TIFF and ZIP archives.

**A PDF with no extractable text is the normal case, not a failure.**
Most scanned documents in an enterprise are exactly that. These parsers
set ``needs_ocr`` and return the pages they found, so the pipeline sends
the document to OCR rather than recording that it says nothing.

**An archive is expanded under limits.** Member count, expansion ratio
and nesting are all capped, because an archive is the one format that
can be a few kilobytes on disk and unbounded in memory.
"""

from __future__ import annotations

import io
import zipfile
from collections.abc import Iterator
from typing import TYPE_CHECKING, Any

from app.documents.parser import (
    DocumentParseError,
    ParsedDocument,
    ParsedPage,
    limit_text,
    register,
)
from app.models.enums import DocumentFormat

if TYPE_CHECKING:  # pragma: no cover -- import-time only for the checker
    from docx.document import Document as DocxDocument

MIN_TEXT_CHARACTERS_PER_PAGE = 16
"""Below this a PDF page's text layer is a page number and a header, not
content. Treating it as content is what makes a scanned document look
like a parsed one."""

MAX_ARCHIVE_MEMBERS = 512
MAX_ARCHIVE_RATIO = 100
"""Uncompressed-to-compressed ratio at which an archive is refused. Real
documents compress to perhaps 10:1; a hundredfold is a zip bomb."""
MAX_ARCHIVE_BYTES = 256 * 1024 * 1024
MAX_ARCHIVE_DEPTH = 2
"""How deep nested archives are followed. One archive inside another is
a real thing people do; three levels is not."""

_SKIPPED_ARCHIVE_SUFFIXES = (".exe", ".dll", ".so", ".dylib", ".bin")


def parse_pdf(data: bytes) -> ParsedDocument:
    """A PDF, page by page.

    Raises:
        DocumentParseError: When the file is not a readable PDF at all --
            as distinct from a readable PDF with no text layer, which
            returns normally with ``needs_ocr`` set.
    """
    from pypdf import PdfReader  # noqa: PLC0415 -- a heavy import, only for PDFs
    from pypdf.errors import PyPdfError  # noqa: PLC0415

    document = ParsedDocument(format=DocumentFormat.PDF)
    try:
        reader = PdfReader(io.BytesIO(data))
    except (PyPdfError, ValueError, OSError) as error:
        raise DocumentParseError(f"This PDF could not be opened: {error}") from error

    if reader.is_encrypted:
        # An empty password opens most "encrypted" PDFs, which are
        # permission-locked rather than secret.
        try:
            opened = reader.decrypt("")
        except (PyPdfError, NotImplementedError) as error:
            raise DocumentParseError(
                f"This PDF is encrypted and could not be opened: {error}"
            ) from error
        if not opened:
            raise DocumentParseError("This PDF is encrypted and needs a password.")
        document.add_warning("The PDF was encrypted and was opened with an empty password.")

    for key, value in (reader.metadata or {}).items():
        if value is not None:
            document.metadata[str(key).lstrip("/")] = str(value)[:512]

    for number, page in enumerate(reader.pages, start=1):
        document.pages.append(_pdf_page(number, page, document))

    if not document.pages:
        raise DocumentParseError("This PDF contains no pages.")

    empty = sum(1 for page in document.pages if not page.has_text_layer)
    if empty:
        document.needs_ocr = True
        document.add_warning(
            f"{empty} of {len(document.pages)} page(s) carry no usable text layer " "and need OCR."
        )
    return document


def _pdf_page(number: int, page: Any, document: ParsedDocument) -> ParsedPage:
    """One PDF page, with a failed extraction recorded rather than raised.

    One unreadable page must not lose the other two hundred, so a failure
    here becomes a warning and an OCR flag on that page alone.
    """
    from pypdf.errors import PyPdfError  # noqa: PLC0415

    try:
        text = page.extract_text() or ""
    except (PyPdfError, KeyError, TypeError, ValueError, UnicodeDecodeError) as error:
        document.add_warning(f"Page {number} could not be read ({error}); it needs OCR.")
        text = ""

    box = getattr(page, "mediabox", None)
    return ParsedPage(
        number=number,
        text=limit_text(text.strip(), document),
        width=float(box.width) if box is not None else None,
        height=float(box.height) if box is not None else None,
        rotation=int(getattr(page, "rotation", 0) or 0),
        has_text_layer=len(text.strip()) >= MIN_TEXT_CHARACTERS_PER_PAGE,
    )


def parse_docx(data: bytes) -> ParsedDocument:
    """A DOCX, with its tables rendered as pipe-delimited rows.

    Tables are emitted inline in document order so the table extractor
    finds them where the reader would. Headers, footers and footnotes are
    not in the body and are recorded as dropped rather than silently
    omitted.
    """
    import docx  # noqa: PLC0415 -- a heavy import, only for DOCX
    from docx.opc.exceptions import PackageNotFoundError  # noqa: PLC0415

    document = ParsedDocument(format=DocumentFormat.DOCX)
    try:
        source = docx.Document(io.BytesIO(data))
    except (
        PackageNotFoundError,
        KeyError,
        ValueError,
        OSError,
        zipfile.BadZipFile,
    ) as error:
        # BadZipFile is not an OSError, so a corrupt DOCX escaped this
        # handler entirely and crashed the worker instead of being recorded
        # as a parse failure on the document.
        raise DocumentParseError(f"This DOCX could not be opened: {error}") from error

    core = source.core_properties
    for name in ("title", "author", "subject", "category", "comments", "last_modified_by"):
        value = getattr(core, name, None)
        if value:
            document.metadata[name] = str(value)[:512]

    blocks = list(_docx_blocks(source))
    text = "\n\n".join(block for block in blocks if block.strip())
    document.pages = [ParsedPage(number=1, text=limit_text(text, document))]

    tables = len(source.tables)
    if tables:
        document.metadata["table_count"] = str(tables)
    if any(section.header or section.footer for section in source.sections):
        document.add_warning("Headers and footers were not included in the extracted text.")
    if not text.strip():
        document.needs_ocr = True
        document.add_warning("The DOCX body holds no text; it may contain only images.")
    return document


def _docx_blocks(source: DocxDocument) -> Iterator[str]:
    """Paragraphs and tables in the order the document lays them out.

    python-docx exposes paragraphs and tables as separate collections, so
    reading them in turn would put every table after every paragraph. The
    body's own XML children preserve the order a reader sees.
    """
    from docx.table import Table  # noqa: PLC0415
    from docx.text.paragraph import Paragraph  # noqa: PLC0415

    body = source.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, source).text
        elif child.tag.endswith("}tbl"):
            table = Table(child, source)
            yield "\n".join(
                "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
                for row in table.rows
            )


def parse_xlsx(data: bytes) -> ParsedDocument:
    """An XLSX workbook, one page per sheet.

    Loaded with ``data_only``, so a cell holding a formula yields its
    cached result rather than ``=SUM(A1:A9)`` -- the value is what a
    reader sees and what an extractor wants. Where a workbook has never
    been opened by Excel that cache is empty, and this says so.
    """
    from openpyxl import load_workbook  # noqa: PLC0415 -- a heavy import
    from openpyxl.utils.exceptions import InvalidFileException  # noqa: PLC0415

    document = ParsedDocument(format=DocumentFormat.XLSX)
    try:
        workbook = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    except (InvalidFileException, KeyError, ValueError, OSError, zipfile.BadZipFile) as error:
        raise DocumentParseError(f"This workbook could not be opened: {error}") from error

    try:
        empty_formula_cache = False
        for number, sheet in enumerate(workbook.worksheets, start=1):
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value) for value in row]
                if any(cell.strip() for cell in cells):
                    rows.append("| " + " | ".join(cells) + " |")
            if not rows and sheet.max_row and sheet.max_row > 1:
                empty_formula_cache = True
            document.pages.append(
                ParsedPage(
                    number=number,
                    text=limit_text(f"# {sheet.title}\n" + "\n".join(rows), document),
                )
            )
        document.metadata["sheet_count"] = str(len(document.pages))
        if empty_formula_cache:
            document.add_warning(
                "A sheet has rows but no cached values; the workbook may hold "
                "formulas that have never been evaluated."
            )
    finally:
        workbook.close()

    if not document.pages:
        raise DocumentParseError("This workbook contains no sheets.")
    return document


def parse_image(data: bytes) -> ParsedDocument:
    """A raster image: geometry only, with OCR left to the OCR engine.

    An image carries no text layer by definition, so this always sets
    ``needs_ocr``. It exists to record the page geometry OCR needs and to
    fail early on a payload that is not a decodable image.
    """
    return _image_document(data, DocumentFormat.IMAGE)


def parse_tiff(data: bytes) -> ParsedDocument:
    """A TIFF, one page per frame.

    TIFF is the multi-page scan format, and collapsing a forty-page fax
    to one page loses every page reference downstream.
    """
    return _image_document(data, DocumentFormat.TIFF)


def _image_document(data: bytes, fmt: DocumentFormat) -> ParsedDocument:
    """Geometry for every frame of an image."""
    from PIL import Image, UnidentifiedImageError  # noqa: PLC0415 -- a heavy import

    document = ParsedDocument(format=fmt, needs_ocr=True)
    try:
        with Image.open(io.BytesIO(data)) as image:
            document.metadata["mode"] = image.mode
            document.metadata["image_format"] = image.format or str(fmt)
            frames = getattr(image, "n_frames", 1)
            for number in range(1, frames + 1):
                if frames > 1:
                    image.seek(number - 1)
                document.pages.append(
                    ParsedPage(
                        number=number,
                        width=float(image.width),
                        height=float(image.height),
                        has_text_layer=False,
                        image_count=1,
                    )
                )
    except (UnidentifiedImageError, OSError, ValueError) as error:
        raise DocumentParseError(f"This image could not be decoded: {error}") from error

    document.metadata["frame_count"] = str(len(document.pages))
    document.add_warning("Images carry no text layer; this document needs OCR to be read.")
    return document


def parse_zip(data: bytes, depth: int = 0) -> ParsedDocument:
    """A ZIP archive, with each member parsed in turn.

    Raises:
        DocumentParseError: When the archive is unreadable, or when it
            declares an expansion beyond the configured limits.
    """
    from app.documents.detection import detect_format  # noqa: PLC0415 -- avoids a cycle
    from app.documents.parser import parser_for  # noqa: PLC0415 -- avoids a cycle

    document = ParsedDocument(format=DocumentFormat.ZIP)
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except (zipfile.BadZipFile, OSError, ValueError) as error:
        raise DocumentParseError(f"This archive could not be opened: {error}") from error

    with archive:
        members = [item for item in archive.infolist() if not item.is_dir()]
        _guard_archive(members, len(data))
        document.metadata["member_count"] = str(len(members))

        for member in members:
            name = member.filename
            if name.lower().endswith(_SKIPPED_ARCHIVE_SUFFIXES):
                document.add_warning(f"{name!r} is an executable and was not parsed.")
                continue
            try:
                payload = archive.read(member)
            except (zipfile.BadZipFile, RuntimeError, OSError) as error:
                document.add_warning(f"{name!r} could not be extracted: {error}")
                continue
            child = _parse_member(payload, name, depth, document, detect_format, parser_for)
            if child is not None:
                child.metadata["archive_member"] = name
                document.attachments.append(child)

    document.pages = [
        ParsedPage(number=index, text=child.text)
        for index, child in enumerate(document.attachments, start=1)
    ]
    document.needs_ocr = any(child.needs_ocr for child in document.attachments)
    if not document.attachments:
        document.add_warning("No member of this archive could be parsed.")
    return document


def _parse_member(
    payload: bytes,
    name: str,
    depth: int,
    document: ParsedDocument,
    detect_format: Any,
    parser_for: Any,
) -> ParsedDocument | None:
    """One archive member, or ``None`` if it could not be parsed."""
    guess = detect_format(payload, filename=name)
    if not guess.is_known:
        document.add_warning(f"{name!r} is of an unrecognised format and was skipped.")
        return None
    if guess.format is DocumentFormat.ZIP:
        if depth >= MAX_ARCHIVE_DEPTH:
            document.add_warning(
                f"{name!r} is a nested archive beyond {MAX_ARCHIVE_DEPTH} levels "
                "and was not expanded."
            )
            return None
        return parse_zip(payload, depth + 1)
    try:
        parsed: ParsedDocument = parser_for(guess.format)(payload)
        return parsed
    except DocumentParseError as error:
        document.add_warning(f"{name!r} could not be parsed: {error}")
        return None


def _guard_archive(members: list[zipfile.ZipInfo], compressed_size: int) -> None:
    """Refuse an archive that would expand beyond the limits.

    Checked against the *declared* sizes before reading anything, which
    is the only point at which a zip bomb can still be refused cheaply.

    Raises:
        DocumentParseError: When any limit would be exceeded.
    """
    if len(members) > MAX_ARCHIVE_MEMBERS:
        raise DocumentParseError(
            f"This archive declares {len(members)} members, above the "
            f"{MAX_ARCHIVE_MEMBERS} limit."
        )
    declared = sum(member.file_size for member in members)
    if declared > MAX_ARCHIVE_BYTES:
        raise DocumentParseError(
            f"This archive declares {declared} bytes uncompressed, above the "
            f"{MAX_ARCHIVE_BYTES} limit."
        )
    if compressed_size and declared / compressed_size > MAX_ARCHIVE_RATIO:
        raise DocumentParseError(
            f"This archive expands {declared / compressed_size:.0f}-fold, above the "
            f"{MAX_ARCHIVE_RATIO}-fold limit; it looks like a zip bomb."
        )


def _register_all() -> None:
    """Wire every binary parser into the registry."""
    register(DocumentFormat.PDF, parse_pdf)
    register(DocumentFormat.DOCX, parse_docx)
    register(DocumentFormat.XLSX, parse_xlsx)
    register(DocumentFormat.IMAGE, parse_image)
    register(DocumentFormat.TIFF, parse_tiff)
    register(DocumentFormat.ZIP, parse_zip)


_register_all()


__all__ = [
    "MAX_ARCHIVE_BYTES",
    "MAX_ARCHIVE_DEPTH",
    "MAX_ARCHIVE_MEMBERS",
    "MAX_ARCHIVE_RATIO",
    "MIN_TEXT_CHARACTERS_PER_PAGE",
    "parse_docx",
    "parse_image",
    "parse_pdf",
    "parse_tiff",
    "parse_xlsx",
    "parse_zip",
]
